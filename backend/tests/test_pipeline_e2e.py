from docx import Document
import io

from app.modules.lessons import service as lesson_service
from app.modules.slides import service as slides_service
from app.modules.topics import service as topic_service


def test_minimal_pipeline_from_docx_upload_to_rendered_slides(client, monkeypatch):
    monkeypatch.setattr(
        topic_service,
        "call_extract_topics",
        lambda sections: [
            {
                "name": "ECG Rate Assessment",
                "summary": "Estimate heart rate from ECG rhythm strips.",
                "tags": ["ecg", "rate"],
            }
        ],
    )
    monkeypatch.setattr(
        lesson_service,
        "call_generate_lesson_plan",
        lambda outline_items, num_lessons, duration_minutes, target_audience: {
            "title": "ECG Basics",
            "lessons": [
                {
                    "title": "ECG Rate Assessment",
                    "objectives": ["Estimate heart rate from an ECG strip"],
                    "timeline": [
                        {
                            "time": "0-10 min",
                            "activity": "Review ECG paper speed and large boxes",
                        }
                    ],
                    "topics": ["ECG Rate Assessment"],
                }
            ],
        },
    )
    monkeypatch.setattr(
        slides_service,
        "call_generate_yaml_slides",
        lambda lesson_title, objectives, topics: (
            'lesson_title: "ECG Rate Assessment"\n'
            "slides:\n"
            '  - title: "Rate check"\n'
            '    type: "content"\n'
            "    bullets:\n"
            '      - "Count large boxes between R waves"\n'
        ),
    )

    version_response = client.post(
        "/curriculum/version",
        json={"name": "ECG Intro", "year": 2026},
    )
    assert version_response.status_code == 200
    version_id = version_response.json()["id"]

    upload_response = client.post(
        f"/documents/{version_id}/upload",
        files={
            "file": (
                "ecg-intro.docx",
                _build_docx_bytes(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    assert upload_response.status_code == 200
    assert upload_response.json()["status"] == "parsed"
    material_id = upload_response.json()["id"]

    sections_response = client.get(f"/documents/{material_id}/sections")
    assert sections_response.status_code == 200
    sections = sections_response.json()
    assert len(sections) == 1
    assert sections[0]["title"] == "ECG Rate Assessment"

    extract_response = client.post(
        "/topics/extract",
        json={
            "curriculum_version_id": version_id,
            "section_ids": [sections[0]["id"]],
        },
    )
    assert extract_response.status_code == 200
    topics = extract_response.json()["topics"]
    assert topics == [{"id": topics[0]["id"], "name": "ECG Rate Assessment"}]

    outline_response = client.post(
        "/curriculum/outline",
        json={
            "curriculum_version_id": version_id,
            "items": [
                {
                    "type": "topic",
                    "topic_id": topics[0]["id"],
                    "title": "ECG Rate Assessment",
                }
            ],
        },
    )
    assert outline_response.status_code == 200
    outline_id = outline_response.json()["id"]

    lesson_response = client.post(
        "/lessons/generate",
        json={
            "outline_id": outline_id,
            "title": "ECG Rate Assessment",
            "num_lessons": 1,
            "duration_minutes": 45,
            "target_audience": "nursing learners",
        },
    )
    assert lesson_response.status_code == 200
    lesson_id = lesson_response.json()["id"]

    slides_response = client.post(f"/slides/generate/{lesson_id}")
    assert slides_response.status_code == 200
    assert slides_response.json()["yaml"].startswith(
        'lesson_title: "ECG Rate Assessment"'
    )

    render_response = client.post(f"/slides/render/{lesson_id}")
    assert render_response.status_code == 200
    assert "Rate check" in render_response.json()["html"]


def _build_docx_bytes() -> bytes:
    document = Document()
    document.add_heading("ECG Rate Assessment", level=1)
    document.add_paragraph("Use large boxes between R waves to estimate rate.")
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()
